from docx import Document

from carnatic_bingo_generator import GRID_LABELS, parse_mcqs, parse_concert_clues
from conftest import CONCERT_DOCX, MCQ_DOCX
from carnatic_bingo_generator import generate_sheet


def _bingo_grid_table(doc):
    """First 6x6 table after info rows is the bingo grid."""
    for table in doc.tables:
        if len(table.rows) == 6 and len(table.columns) == 6:
            return table
    raise AssertionError("Bingo grid table not found")


def _i_heard_labels(doc):
    labels = []
    for table in doc.tables:
        if len(table.columns) == 2 and len(table.rows) == 12:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text[0] in "ABCDE" and "." in text[:4]:
                        labels.append(text.split(".")[0])
    return labels


def test_generate_docx_structure(env_isolated, tmp_state_file, tmp_output_dir, bingo_root):
    mcqs = parse_mcqs(str(MCQ_DOCX))
    entry = generate_sheet(state_path=tmp_state_file, output_dir=tmp_output_dir)
    doc = Document(entry["path"])

    subtitles = [p.text for p in doc.paragraphs if "Bingo Sheet #" in p.text]
    assert any("Bingo Sheet #1" in t for t in subtitles)

    grid = _bingo_grid_table(doc)
    a1 = grid.rows[1].cells[1].text.strip()
    assert a1 == ""
    c3 = grid.rows[3].cells[3].text
    assert "★ Answer MCQ ★" in c3

    mcq_snippet = mcqs[0]["question"][:40].replace("\xa0", " ")
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += "\n" + cell.text
    assert mcq_snippet in all_text.replace("\xa0", " ")

    labels = _i_heard_labels(doc)
    assert len(labels) == 24
    assert "C3" not in labels
    assert set(labels) == set(GRID_LABELS)
