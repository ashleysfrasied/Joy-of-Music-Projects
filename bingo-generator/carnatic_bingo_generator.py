"""
Carnatic Bingo Sheet Generator
Run: python3 carnatic_bingo_generator.py
"""

from __future__ import annotations

import json
import os
import random
import re
from datetime import datetime, timezone
from typing import Any

from docx import Document

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

CONCERT_LIST_FILENAME = "input_bingo_questions.docx"
MCQ_LIST_FILENAME = "input_multiple_choice_Q&A.docx"
TEMPLATE_FILENAME = "input_reference_bingo_sheet.docx"
LOGO_FILENAME = "BrandLogoTJOMWords-SideBySide-WhiteBG.jpg"
LOGO_PNG_FILENAME = "logo.png"
OUTPUT_SUBDIR = "completed-bingo-sheets"
CONCERT_TITLE_SKIP = "Bingo quiz concerts"
EXPECTED_CLUES = 24
EXPECTED_MCQS = 20
SHUFFLE_MAX_RETRIES = 100

GRID_LABELS = [
    f"{r}{c}" for r in "ABCDE" for c in range(1, 6) if not (r == "C" and c == 3)
]

_SHEET_TITLE_RE = re.compile(
    r"(Carnatic Bingo\s*\|\s*Bingo Sheet #)\d+", re.IGNORECASE
)
_CLUE_LABEL_RE = re.compile(r"^([A-E][1-5])\.\s*", re.IGNORECASE)

_OPTION_RE = re.compile(r"^[A-D]\.\s*(.+)$", re.IGNORECASE)
_ANSWER_RE = re.compile(r"^Answer:\s*([A-D])\s*$", re.IGNORECASE)


def get_bingo_root() -> str:
    return os.environ.get("BINGO_ROOT", SCRIPT_DIR)


def get_state_path() -> str:
    return os.environ.get("BINGO_STATE_FILE", os.path.join(get_bingo_root(), "bingo_state.json"))


def get_output_dir() -> str:
    return os.environ.get(
        "BINGO_OUTPUT_DIR", os.path.join(get_bingo_root(), OUTPUT_SUBDIR)
    )


def get_template_path(root: str | None = None) -> str:
    bingo_root = root or get_bingo_root()
    path = os.environ.get(
        "BINGO_TEMPLATE_FILE", os.path.join(bingo_root, TEMPLATE_FILENAME)
    )
    if not os.path.isfile(path):
        raise FileNotFoundError(f"Bingo template not found: {path}")
    return path


def _normalize_text(text: str) -> str:
    return text.replace("\xa0", " ").strip()


def parse_concert_clues(path: str) -> list[str]:
    doc = Document(path)
    clues: list[str] = []
    for p in doc.paragraphs:
        text = _normalize_text(p.text)
        if not text or text == CONCERT_TITLE_SKIP:
            continue
        clues.append(text)
    if len(clues) != EXPECTED_CLUES:
        raise ValueError(
            f"Expected {EXPECTED_CLUES} concert clues in {path}, found {len(clues)}"
        )
    return clues


def _parse_mcq_block(block: str) -> dict[str, Any]:
    lines = [_normalize_text(line) for line in block.split("\n") if _normalize_text(line)]
    if not lines:
        raise ValueError("Empty MCQ block")

    answer = None
    options: list[str] = []
    question_lines: list[str] = []

    for line in lines:
        m_ans = _ANSWER_RE.match(line)
        if m_ans:
            answer = m_ans.group(1).upper()
            continue
        m_opt = _OPTION_RE.match(line)
        if m_opt:
            options.append(m_opt.group(1).strip())
            continue
        if not options and answer is None:
            question_lines.append(line)

    question = " ".join(question_lines).strip()
    if not question or len(options) != 4 or answer not in "ABCD":
        raise ValueError(f"Invalid MCQ block: {block[:80]}...")
    return {"question": question, "options": options, "answer": answer}


def parse_mcqs(path: str) -> list[dict[str, Any]]:
    doc = Document(path)
    blocks: list[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text or text.startswith("Bingo Sheet") or "Form" in text:
            continue
        blocks.append(text)

    mcqs = [_parse_mcq_block(b) for b in blocks]
    if len(mcqs) != EXPECTED_MCQS:
        raise ValueError(f"Expected {EXPECTED_MCQS} MCQs in {path}, found {len(mcqs)}")
    return mcqs


def default_state() -> dict[str, Any]:
    return {
        "sheet_number": 0,
        "next_mcq_index": 0,
        "past_clue_orders": [],
        "generated_sheets": [],
    }


def load_state(state_path: str | None = None) -> dict[str, Any]:
    path = state_path or get_state_path()
    if os.path.exists(path):
        with open(path, encoding="utf-8") as f:
            data = json.load(f)
        for key, val in default_state().items():
            data.setdefault(key, val)
        return data
    return default_state()


def save_state(state: dict[str, Any], state_path: str | None = None) -> None:
    path = state_path or get_state_path()
    os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(state, f, indent=2)


def _unique_clue_order(past_orders: list[list[int]], num_clues: int) -> list[int]:
    past_set = {tuple(o) for o in past_orders}
    for _ in range(SHUFFLE_MAX_RETRIES):
        order = list(range(num_clues))
        random.shuffle(order)
        if tuple(order) not in past_set:
            return order
    raise RuntimeError(
        f"Could not find a new clue permutation after {SHUFFLE_MAX_RETRIES} retries"
    )


def next_generation_params(
    num_clues: int = EXPECTED_CLUES,
    num_mcqs: int = EXPECTED_MCQS,
    state_path: str | None = None,
) -> tuple[int, list[int], int, dict[str, Any]]:
    """Advance state and return (sheet_number, clue_order, mcq_index, updated_state)."""
    state = load_state(state_path)
    clue_order = _unique_clue_order(state["past_clue_orders"], num_clues)
    sheet_number = state["sheet_number"] + 1
    mcq_index = state["next_mcq_index"] % num_mcqs

    state["sheet_number"] = sheet_number
    state["past_clue_orders"].append(clue_order)
    state["next_mcq_index"] = (mcq_index + 1) % num_mcqs

    return sheet_number, clue_order, mcq_index, state


def output_filename(sheet_number: int) -> str:
    return f"Carnatic_Bingo_Sheet_{sheet_number:03d}.docx"


def _find_table(doc: Document, rows: int, cols: int):
    for table in doc.tables:
        if len(table.rows) == rows and len(table.columns) == cols:
            return table
    raise ValueError(f"Template missing {rows}x{cols} table")


def _cell_label(cell) -> str:
    text = _normalize_text(cell.text)
    m = _CLUE_LABEL_RE.match(text)
    if m:
        return m.group(1).upper()
    if cell.paragraphs[0].runs:
        run_text = _normalize_text(cell.paragraphs[0].runs[0].text)
        m = _CLUE_LABEL_RE.match(run_text)
        if m:
            return m.group(1).upper()
    raise ValueError(f"Could not parse grid label from clue cell: {text!r}")


def _set_run_text(run, text: str) -> None:
    run.text = text


def _update_sheet_titles(doc: Document, sheet_number: int) -> None:
    replacement = rf"\g<1>{sheet_number}"
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if "Bingo Sheet #" in run.text:
                run.text = _SHEET_TITLE_RE.sub(replacement, run.text)


def _clear_bingo_data_cells(grid_table) -> None:
    for ri in range(1, 6):
        for ci in range(1, 6):
            if ri == 3 and ci == 3:
                continue
            cell = grid_table.rows[ri].cells[ci]
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.text = ""


def _update_i_heard_table(doc: Document, clue_map: dict[str, str]) -> None:
    heard = _find_table(doc, 12, 2)
    for row in heard.rows:
        for cell in row.cells:
            label = _cell_label(cell)
            clue_text = clue_map[label]
            paragraph = cell.paragraphs[0]
            if len(paragraph.runs) >= 2:
                _set_run_text(paragraph.runs[0], f"{label}.  ")
                _set_run_text(paragraph.runs[1], clue_text)
            else:
                paragraph.clear()
                bold_run = paragraph.add_run(f"{label}.  ")
                bold_run.bold = True
                paragraph.add_run(clue_text)


def _update_mcq_cell(cell, mcq: dict[str, Any]) -> None:
    paragraphs = cell.paragraphs
    if len(paragraphs) < 5:
        raise ValueError(
            f"Template MCQ cell expected 5 paragraphs, found {len(paragraphs)}"
        )
    _set_run_text(paragraphs[0].runs[0], mcq["question"])
    for idx, opt in enumerate(mcq["options"]):
        letter = chr(65 + idx)
        _set_run_text(paragraphs[idx + 1].runs[0], f"  \u25a1  {letter})  {opt}")


def build_bingo_docx(
    sheet_number: int,
    clues: list[str],
    clue_order: list[int],
    mcq: dict[str, Any],
    output_path: str,
    template_path: str,
) -> str:
    """Fill template with dynamic sheet number, clues, and MCQ."""
    clue_map = {
        GRID_LABELS[i]: clues[clue_order[i]] for i in range(len(GRID_LABELS))
    }

    doc = Document(template_path)
    _update_sheet_titles(doc, sheet_number)
    _clear_bingo_data_cells(_find_table(doc, 6, 6))
    _update_i_heard_table(doc, clue_map)
    _update_mcq_cell(_find_table(doc, 1, 1).rows[0].cells[0], mcq)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_sheet(
    bingo_root: str | None = None,
    state_path: str | None = None,
    output_dir: str | None = None,
) -> dict[str, Any]:
    """Generate one bingo sheet; update state; return manifest entry."""
    root = bingo_root or get_bingo_root()
    state_file = state_path or get_state_path()
    out_dir = output_dir or get_output_dir()
    os.makedirs(out_dir, exist_ok=True)

    concert_path = os.path.join(root, CONCERT_LIST_FILENAME)
    mcq_path = os.path.join(root, MCQ_LIST_FILENAME)
    template_path = get_template_path(root)

    clues = parse_concert_clues(concert_path)
    mcqs = parse_mcqs(mcq_path)

    sheet_number, clue_order, mcq_index, state = next_generation_params(
        num_clues=len(clues),
        num_mcqs=len(mcqs),
        state_path=state_file,
    )

    filename = output_filename(sheet_number)
    output_path = os.path.join(out_dir, filename)

    build_bingo_docx(
        sheet_number=sheet_number,
        clues=clues,
        clue_order=clue_order,
        mcq=mcqs[mcq_index],
        output_path=output_path,
        template_path=template_path,
    )

    entry = {
        "sheet_number": sheet_number,
        "filename": filename,
        "path": output_path,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "mcq_index": mcq_index,
        "mcq_question": mcqs[mcq_index]["question"],
        "clue_order": clue_order,
    }
    state["generated_sheets"].append(entry)
    save_state(state, state_file)
    return entry


def main() -> None:
    entry = generate_sheet()
    print(f"Generated : {entry['filename']}")
    print(f"Sheet #   : {entry['sheet_number']}")
    print(f"MCQ used  : #{entry['mcq_index'] + 1} — {entry['mcq_question'][:65]}...")
    print(f"Saved to  : {entry['path']}")


if __name__ == "__main__":
    main()
