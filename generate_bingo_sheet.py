#!/usr/bin/env python3
"""
Build a Carnatic Bingo Word document: reads 25 prompts from the outline .docx,
shuffles which clue appears with each square (A1–E5), leaves the play grid blank
for marks, and writes output/Bingo sheet-{n}.docx.

Two pages: sheet info + grid + clue list on page 1; notes and concert-details tables on page 2.

Before writing, all prior output/Bingo sheet-*.docx files are removed so only
the new sheet remains (sheet number still increments from what existed before
cleanup).

Usage:
  python3 generate_bingo_sheet.py
  python3 generate_bingo_sheet.py --source path/to/outline.docx --outdir path/to/output
"""

from __future__ import annotations

import argparse
import random
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table


QUESTION_LINE = re.compile(r"^[A-E][1-5]\.\s*(.+)$", re.IGNORECASE)
DEFAULT_SOURCE = Path(__file__).resolve().parent / "CarnaticBingoSheetBasicOutline.docx"
DEFAULT_OUTDIR = Path(__file__).resolve().parent / "output"

ALL_CELL_CODES = [f"{L}{n}" for L in "ABCDE" for n in range(1, 6)]


def extract_questions(source: Path) -> list[str]:
    doc = Document(source)
    found: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        m = QUESTION_LINE.match(text)
        if m:
            found.append(m.group(1).strip())
    if len(found) != 25:
        raise ValueError(
            f"Expected 25 bingo prompts (A1–E5) in {source}, found {len(found)}."
        )
    return found


def next_sheet_number(outdir: Path) -> int:
    outdir.mkdir(parents=True, exist_ok=True)
    best = 0
    for p in outdir.glob("Bingo sheet-*.docx"):
        m = re.match(r"Bingo sheet-(\d+)\.docx$", p.name, re.IGNORECASE)
        if m:
            best = max(best, int(m.group(1)))
    return best + 1


def remove_old_bingo_sheets(outdir: Path) -> int:
    """Delete every Bingo sheet-*.docx in outdir. Returns how many files removed."""
    outdir.mkdir(parents=True, exist_ok=True)
    removed = 0
    for p in outdir.glob("Bingo sheet-*.docx"):
        if p.is_file():
            p.unlink()
            removed += 1
    return removed


def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size_pt: float = 9,
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.bold = bold


def _set_equal_col_widths(table: Table, widths_inches: list[float]) -> None:
    for row in table.rows:
        for idx, w in enumerate(widths_inches):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(w)


def _uniform_row_heights(table: Table, *, first_row_twips: int = 300, body_twips: int = 300) -> None:
    for ri, row in enumerate(table.rows):
        h = first_row_twips if ri == 0 else body_twips
        tr = row._tr
        tr_pr = tr.get_or_add_trPr()
        tr_h = OxmlElement("w:trHeight")
        tr_h.set(qn("w:val"), str(h))
        tr_h.set(qn("w:hRule"), "atLeast")
        tr_pr.append(tr_h)


def _empty_bingo_grid(table: Table) -> None:
    """6×6: headers only; inner 5×5 cells blank for player marks."""
    _set_cell_text(table.rows[0].cells[0], "", size_pt=8)
    for c in range(5):
        _set_cell_text(table.rows[0].cells[c + 1], str(c + 1), bold=True, size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    letters = "ABCDE"
    for r in range(5):
        _set_cell_text(
            table.rows[r + 1].cells[0],
            letters[r],
            bold=True,
            size_pt=9,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        for c in range(5):
            _set_cell_text(table.rows[r + 1].cells[c + 1], "", size_pt=9)

    total = 6.55
    corner = 0.42
    each = (total - corner) / 5.0
    _set_equal_col_widths(table, [corner] + [each] * 5)
    _uniform_row_heights(table)


def _add_contact_lines(doc: Document) -> None:
    """Labeled lines for handwriting — no table."""
    labels = [
        "Student Name:",
        "Parent Name:",
        "Parent Email Address:",
        "Parent Whatsapp:",
        "Teacher:",
        "Location:",
    ]
    for lab in labels:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(lab)
        run.bold = True
        run.font.size = Pt(10)


def _add_square_clues_section(doc: Document, questions_shuffled: list[str]) -> None:
    """Standard header + one line per square clue."""
    h = doc.add_paragraph()
    h.paragraph_format.space_after = Pt(4)
    hr = h.add_run("Clues for each square")
    hr.bold = True
    hr.font.size = Pt(11)
    for code, clue in zip(ALL_CELL_CODES, questions_shuffled):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(f"{code}. {clue}")
        r.font.size = Pt(8)


def build_sheet(
    questions_shuffled: list[str],
    sheet_number: int,
    out_path: Path,
) -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("The Joy of Music")
    tr.bold = True
    tr.font.size = Pt(14)
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = st.add_run(f"Carnatic Bingo — Bingo sheet-{sheet_number}")
    sr.bold = True
    sr.font.size = Pt(12)

    doc.add_paragraph()

    _add_contact_lines(doc)

    doc.add_paragraph()
    instr = doc.add_paragraph()
    instr.paragraph_format.space_after = Pt(2)
    instr.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    ir = instr.add_run(
        "How to play: at a concert, when you hear or see something that matches one of the hints below, "
        "put a check or X in the matching empty square (A1–E5); which clue goes with each square "
        "reshuffles every time you print a new sheet."
    )
    ir.font.size = Pt(9)
    ir.italic = True

    gh = doc.add_paragraph()
    gh.paragraph_format.space_after = Pt(2)
    ghr = gh.add_run("Bingo grid")
    ghr.bold = True
    ghr.font.size = Pt(10)

    grid = doc.add_table(rows=6, cols=6)
    grid.style = "Table Grid"
    _empty_bingo_grid(grid)
    _uniform_row_heights(grid, first_row_twips=260, body_twips=260)

    doc.add_paragraph()
    _add_square_clues_section(doc, questions_shuffled)

    doc.add_page_break()

    log_data_rows = 10
    log = doc.add_table(rows=1 + log_data_rows, cols=8)
    log.style = "Table Grid"
    headers = [
        "Bingo Cell",
        "Song",
        "Ragam",
        "Thalam",
        "Composer",
        "Improv (A/T/N/K)",
        "Composition Type",
        "Concert",
    ]
    # ~usable text width for US Letter minus margins (8.5" − 1.1")
    usable_in = 7.4
    col8 = usable_in / 8.0
    col7 = usable_in / 7.0
    for j, h in enumerate(headers):
        _set_cell_text(log.rows[0].cells[j], h, bold=True, size_pt=8)
    for r in range(1, 1 + log_data_rows):
        for j in range(8):
            _set_cell_text(log.rows[r].cells[j], "", size_pt=10)
    _set_equal_col_widths(log, [col8] * 8)
    _uniform_row_heights(log, first_row_twips=360, body_twips=520)

    gap = doc.add_paragraph()
    gap.paragraph_format.space_after = Pt(6)

    cdh = doc.add_paragraph()
    cdh.paragraph_format.space_after = Pt(2)
    rcd = cdh.add_run("Concert details")
    rcd.bold = True
    rcd.font.size = Pt(9)

    det_data_rows = 10
    det = doc.add_table(rows=1 + det_data_rows, cols=7)
    det.style = "Table Grid"
    top = [
        "Concert #",
        "Date",
        "Artist",
        "Accompanists",
        "Notes (something you noticed)",
        "Venue",
        "Location",
    ]
    w7 = col7
    for j, h in enumerate(top):
        _set_cell_text(det.rows[0].cells[j], h, bold=True, size_pt=8)
    for r in range(1, 1 + det_data_rows):
        for j in range(7):
            _set_cell_text(det.rows[r].cells[j], "", size_pt=10)
    _set_equal_col_widths(det, [w7] * 7)
    _uniform_row_heights(det, first_row_twips=360, body_twips=520)

    doc.save(out_path)


def main() -> None:
    ap = argparse.ArgumentParser(description="Generate a randomized Carnatic Bingo sheet.")
    ap.add_argument("--source", type=Path, default=DEFAULT_SOURCE, help="Outline .docx with A1–E5 lines")
    ap.add_argument("--outdir", type=Path, default=DEFAULT_OUTDIR, help="Output directory")
    ap.add_argument("--number", type=int, default=None, help="Force sheet number")
    args = ap.parse_args()

    source: Path = args.source.expanduser().resolve()
    outdir: Path = args.outdir.expanduser().resolve()
    if not source.is_file():
        raise SystemExit(f"Source document not found: {source}")

    questions = extract_questions(source)
    shuffled = questions.copy()
    random.shuffle(shuffled)

    num = args.number if args.number is not None else next_sheet_number(outdir)
    removed = remove_old_bingo_sheets(outdir)
    out_path = outdir / f"Bingo sheet-{num}.docx"

    build_sheet(shuffled, num, out_path)
    print(f"Wrote {out_path} (removed {removed} prior Bingo sheet file(s))")


if __name__ == "__main__":
    main()
